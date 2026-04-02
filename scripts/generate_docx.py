from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

BASE = os.path.dirname(os.path.abspath(__file__))
LOGO = os.path.join(BASE, "kordas_logo_final.png")

# Colors
NAVY = RGBColor(15, 23, 42)
BLUE = RGBColor(59, 130, 246)
GRAY = RGBColor(100, 116, 139)
LIGHT_BG = RGBColor(248, 250, 252)
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = tcBorders.makeelement(qn(f'w:{edge}'), {
                qn('w:val'): val.get('val', 'single'),
                qn('w:sz'): val.get('sz', '4'),
                qn('w:color'): val.get('color', '000000'),
                qn('w:space'): '0'
            })
            tcBorders.append(el)
    tcPr.append(tcBorders)

def style_header_cell(cell, text):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = WHITE
    run.font.name = 'Segoe UI'
    set_cell_shading(cell, '0F172A')
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def style_data_cell(cell, text, bold=False, color=None, bg=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    run.font.color.rgb = color or NAVY
    run.font.name = 'Segoe UI'
    if bg:
        set_cell_shading(cell, bg)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = NAVY
        run.font.name = 'Segoe UI'
        # Blue underline
        border_p = doc.add_paragraph()
        border_p.paragraph_format.space_before = Pt(0)
        border_p.paragraph_format.space_after = Pt(12)
        pPr = border_p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        bottom = pBdr.makeelement(qn('w:top'), {
            qn('w:val'): 'single', qn('w:sz'): '8',
            qn('w:color'): '3B82F6', qn('w:space'): '1'
        })
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = BLUE
        run.font.name = 'Segoe UI'
    elif level == 3:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
        run.font.name = 'Segoe UI'
    return p

def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(10)
    run.font.color.rgb = color or NAVY
    run.font.name = 'Segoe UI'
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = NAVY
        run.font.name = 'Segoe UI'
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = NAVY
        run.font.name = 'Segoe UI'
    else:
        p.text = ''
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = NAVY
        run.font.name = 'Segoe UI'
    return p

def create_scope_table(doc, rows):
    table = doc.add_table(rows=len(rows)+1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(['#', 'Zadanie', 'Opis']):
        style_header_cell(table.rows[0].cells[i], h)
    # Data
    for r, (num, task, desc) in enumerate(rows, 1):
        alt_bg = 'F8FAFC' if r % 2 == 0 else None
        style_data_cell(table.rows[r].cells[0], num, bold=True, bg=alt_bg)
        style_data_cell(table.rows[r].cells[1], task, bold=True, bg=alt_bg)
        style_data_cell(table.rows[r].cells[2], desc, bg=alt_bg)
    # Column widths
    for row in table.rows:
        row.cells[0].width = Cm(1)
        row.cells[1].width = Cm(4)
        row.cells[2].width = Cm(12)
    return table

def create_comparison_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        style_header_cell(table.rows[0].cells[i], h)
    for r, row_data in enumerate(rows, 1):
        alt_bg = 'F8FAFC' if r % 2 == 0 else None
        for c, val in enumerate(row_data):
            is_price = 'zł' in val or val.startswith('**')
            clean_val = val.replace('**', '')
            style_data_cell(table.rows[r].cells[c], clean_val, bold=is_price, color=BLUE if is_price else None, bg=alt_bg)
    return table

# ============================================================
# BUILD DOCUMENT
# ============================================================

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default font
style = doc.styles['Normal']
font = style.font
font.name = 'Segoe UI'
font.size = Pt(10)
font.color.rgb = NAVY

# ---- HEADER WITH LOGO ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
if os.path.exists(LOGO):
    run = p.add_run()
    run.add_picture(LOGO, width=Inches(3.5))

# ---- TITLE ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(20)
run = p.add_run('OFERTA CENOWA')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = NAVY
run.font.name = 'Segoe UI'

p = doc.add_paragraph()
run = p.add_run('Projekt „Antyszpak" — Rewizja 2.0')
run.font.size = Pt(14)
run.font.color.rgb = GRAY
run.font.name = 'Segoe UI'

p = doc.add_paragraph()
run = p.add_run('Kod oferty: KORDAS/AI/2026/04  |  Data: 01.04.2026')
run.font.size = Pt(10)
run.font.color.rgb = GRAY
run.font.name = 'Segoe UI'

# Blue divider
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(16)
pPr = p._element.get_or_add_pPr()
pBdr = pPr.makeelement(qn('w:pBdr'), {})
bottom = pBdr.makeelement(qn('w:bottom'), {
    qn('w:val'): 'single', qn('w:sz'): '12',
    qn('w:color'): '3B82F6', qn('w:space'): '1'
})
pBdr.append(bottom)
pPr.append(pBdr)

# ---- PARTIES ----
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

cell_l = table.rows[0].cells[0]
cell_l.text = ''
p = cell_l.paragraphs[0]
run = p.add_run('OFERENT\n')
run.bold = True; run.font.size = Pt(8); run.font.color.rgb = BLUE; run.font.name = 'Segoe UI'
run = p.add_run('Marcin Kordas\n')
run.bold = True; run.font.size = Pt(11); run.font.color.rgb = NAVY; run.font.name = 'Segoe UI'
for line in ['NIP: 6772446064', 'Przybyszewskiego 30/2, 30-128 Kraków', 'biuro@kordas.tech | +48 797 252 208']:
    run = p.add_run(line + '\n')
    run.font.size = Pt(9); run.font.color.rgb = GRAY; run.font.name = 'Segoe UI'

cell_r = table.rows[0].cells[1]
cell_r.text = ''
p = cell_r.paragraphs[0]
run = p.add_run('ODBIORCA\n')
run.bold = True; run.font.size = Pt(8); run.font.color.rgb = BLUE; run.font.name = 'Segoe UI'
run = p.add_run('AGROBOTYKA sp. z o.o.\n')
run.bold = True; run.font.size = Pt(11); run.font.color.rgb = NAVY; run.font.name = 'Segoe UI'
for line in ['NIP: 9462751716', 'ul. Gospodarcza 26, 20-213 Lublin']:
    run = p.add_run(line + '\n')
    run.font.size = Pt(9); run.font.color.rgb = GRAY; run.font.name = 'Segoe UI'

# ---- UWAGA WSTĘPNA ----
add_heading_styled(doc, 'Uwaga wstępna', 1)
add_body(doc, 'Niniejsza oferta stanowi odpowiedź na zapytania ofertowe z dnia 01.04.2026 dotyczące usług „Z1 — Rozwój algorytmu AI do detekcji obrazu" oraz „Z2 — Trenowanie algorytmu detekcji w warunkach terenowych" w ramach projektu „Antyszpak".')
add_body(doc, 'Oferta aktualizuje wyceny KORDAS/AI/2025/09/A1 oraz B1 z dnia 30.09.2025, których 30-dniowy okres ważności upłynął 30.10.2025. Aktualizacja wynika z:')
add_bullet(doc, 'upływu czasu i zmiany warunków rynkowych w segmencie usług AI/ML,')
add_bullet(doc, 'istotnych zmian w zakresie prac (twarde progi dokładności 65%/75%, trening na danych Zamawiającego, integracja z prototypem, praca ciągła IoT),')
add_bullet(doc, 'zmiany harmonogramu realizacji.')
add_body(doc, 'Dla każdej usługi przedstawiono dwa warianty — Bazowy i Profesjonalny — różniące się zakresem, gwarancjami i ceną.')
add_body(doc, 'Zastrzeżenie: Niniejsza oferta opiera się na dokumentach Z1 i Z2 dostarczonych 01.04.2026. W przypadku dalszych zmian w wymaganiach — wycena podlega rewizji.', italic=True, color=GRAY)

# ======================================================================
# Z1
# ======================================================================
add_heading_styled(doc, 'USŁUGA Z1 — Rozwój algorytmu AI do detekcji obrazu', 1)
add_body(doc, 'Termin realizacji wg zapytania: listopad 2026 – marzec 2027 (5 miesięcy)', italic=True, color=GRAY)

# Z1-A
add_heading_styled(doc, 'Wariant Z1-A — Bazowy (Proof-of-Concept)', 2)
add_body(doc, 'Kod oferty: KORDAS/AI/2026/04/Z1-BASE', color=GRAY)
add_body(doc, 'Przedmiot: Dostarczenie działającego prototypu algorytmu detekcji ptaków na RPi5, z walidacją w symulowanych warunkach operacyjnych.')

add_heading_styled(doc, 'Zakres prac:', 3)
create_scope_table(doc, [
    ('1', 'Dobór modelu', 'Selekcja i konfiguracja YOLO (YOLOv8n/v10n) zoptymalizowanego pod RPi5. Klasa: „bird".'),
    ('2', 'Optymalizacja', 'Konwersja do TFLite/ONNX/NCNN dla ARM. Best-effort czas reakcji.'),
    ('3', 'Trening', 'Fine-tuning na danych publicznych (COCO, iNaturalist, CUB-200).'),
    ('4', 'Punkt celowania', 'Środek geometryczny bounding-box.'),
    ('5', 'API', 'REST API: współrzędne (x, y), confidence, timestamp.'),
    ('6', 'Walidacja', '3 scenariusze symulacyjne. Raport: Precision, Recall, RMSE.'),
    ('7', 'Dostawa', 'Kod źródłowy, wagi modelu, skrypt uruchomieniowy, README.'),
])

add_heading_styled(doc, 'Wyłączenia z zakresu Z1:', 3)
exclusions_z1a = [
    'Trening na danych Zamawiającego (nagrania z sadów Agrobotyki)',
    'Gwarancja progu dokładności 65%',
    'Redukcja false positives specyficznych dla środowiska sadowego',
    'Integracja z prototypem urządzenia',
    'Przygotowanie i oznaczanie zbiorów danych Zamawiającego',
    'Rozszerzona dokumentacja techniczna',
]
for e in exclusions_z1a:
    add_bullet(doc, e)

add_heading_styled(doc, 'Warunki:', 3)
p = add_body(doc, '')
p.text = ''
for label, val in [('Nakład pracy: ', '80–100 roboczogodzin'), ('Cena ryczałtowa: ', '27 060,00 zł brutto (22 000,00 zł netto + 23% VAT)'), ('Czas realizacji: ', '8 tygodni od podpisania umowy'), ('Płatności: ', '50% zaliczka + 50% po dostawie')]:
    run = p.add_run(label)
    run.bold = True; run.font.size = Pt(10); run.font.color.rgb = NAVY; run.font.name = 'Segoe UI'
    run = p.add_run(val + '\n')
    run.font.size = Pt(10); run.font.color.rgb = NAVY; run.font.name = 'Segoe UI'

# Z1-B
doc.add_page_break()
add_heading_styled(doc, 'Wariant Z1-B — Profesjonalny (Pełna realizacja Z1)', 2)
add_body(doc, 'Kod oferty: KORDAS/AI/2026/04/Z1-PRO', color=GRAY)
add_body(doc, 'Przedmiot: Kompleksowa realizacja zakresu Z1 zgodnie z opisem usługi, włącznie z treningiem na danych Zamawiającego, integracją z prototypem oraz gwarancją progów.')

add_heading_styled(doc, 'Zakres prac:', 3)
create_scope_table(doc, [
    ('1', 'Specyfikacja', 'Warsztat z Zamawiającym. Aneks Specyfikacyjny: parametry kamery, warunki sceny, protokół komunikacji.'),
    ('2', 'Pipeline danych', 'Analiza, czyszczenie, augmentacja zbiorów — publicznych + Agrobotyka. Wsparcie przy anotacji.'),
    ('3', 'Trening modelu', 'YOLO/DETR, trening na danych łączonych, walidacja krzyżowa.'),
    ('4', 'Optymalizacja RPi5', 'Kwantyzacja, profilowanie, testy termiczne. Cel: < 1 s end-to-end.'),
    ('5', 'Redukcja FP', 'Multi-frame confirmation, confidence thresholding, logika kontekstowa.'),
    ('6', 'Punkt celowania', 'Konfigurowalna strategia (bbox / centroid / ważony).'),
    ('7', 'API produkcyjne', 'REST + gRPC. Dokumentacja OpenAPI. Health-check.'),
    ('8', 'Integracja', 'Komunikacja z prototypem (MQTT/serial/HTTP). Wymiana danych.'),
    ('9', 'Walidacja', '7+ scenariuszy. Raport per scenariusz. Cel: ≥ 65% dokładności.'),
    ('10', 'Dokumentacja', 'Pełna techniczna + transfer IP + kod + modele.'),
])

add_body(doc, 'Uwaga dot. progu 65%: Osiągnięcie gwarantowane pod warunkiem dostarczenia min. 500 anotowanych klatek (≥ 720p) z ptakami w warunkach sadowych. Przy braku — tryb best-effort.', italic=True, color=GRAY)

add_heading_styled(doc, 'Warunki:', 3)
p = add_body(doc, '')
p.text = ''
for label, val in [
    ('Nakład pracy: ', '220–280 roboczogodzin'),
    ('Cena ryczałtowa: ', '73 800,00 zł brutto (60 000,00 zł netto + 23% VAT)'),
    ('Czas realizacji: ', '18 tygodni od Aneksu Specyfikacyjnego'),
    ('Płatności: ', '30% zaliczka → 30% po MVP → 30% po odbiorze → 10% stabilizacja (30 dni)')
]:
    run = p.add_run(label)
    run.bold = True; run.font.size = Pt(10); run.font.color.rgb = NAVY; run.font.name = 'Segoe UI'
    run = p.add_run(val + '\n')
    run.font.size = Pt(10); run.font.color.rgb = NAVY; run.font.name = 'Segoe UI'

# ======================================================================
# Z2
# ======================================================================
doc.add_page_break()
add_heading_styled(doc, 'USŁUGA Z2 — Trenowanie algorytmu w warunkach terenowych', 1)
add_body(doc, 'Termin realizacji wg zapytania: kwiecień 2027 – lipiec 2027 (4 miesiące)', italic=True, color=GRAY)
add_body(doc, 'Uwaga: Z2 jest kontynuacją Z1 (Faza II). Wymaga uprzedniego zakończenia Z1 w wariancie Profesjonalnym (Z1-B). Realizacja Z2 po Z1-A wymaga dodatkowych prac — wycena na zapytanie.', italic=True, color=GRAY)

# Z2-A
add_heading_styled(doc, 'Wariant Z2-A — Bazowy (Adaptacja dzienna)', 2)
add_body(doc, 'Kod oferty: KORDAS/AI/2026/04/Z2-BASE', color=GRAY)

add_heading_styled(doc, 'Zakres prac:', 3)
create_scope_table(doc, [
    ('1', 'Analiza pilotaży', 'Analiza materiałów z instalacji pilotażowych. Identyfikacja wyzwań.'),
    ('2', 'Stabilizacja tła', 'Filtrowanie: ruch roślinności, cienie, obiekty ruchome.'),
    ('3', 'Retraining', 'Trening na danych z pilotaży (dzień, 3 warianty pogodowe).'),
    ('4', 'Benchmarking', 'Testy porównawcze na danych pilotażowych. Raport.'),
    ('5', 'Optymalizacja IoT', 'Praca ciągła na RPi5: zarządzanie pamięcią, basic watchdog.'),
    ('6', 'API streaming', 'Praca online (strumień wideo).'),
    ('7', 'Dostawa', 'Zaktualizowany kod, model, instrukcja wdrożenia.'),
])

add_heading_styled(doc, 'Wyłączenia:', 3)
for e in ['Gwarancja progu 75%', 'Praca nocna / słabe oświetlenie', 'Warunki ekstremalne (mgła, śnieg, burza)', 'Integracja z aplikacją użytkownika', 'Fizyczne testy terenowe (testy na danych — tak)', 'Pełna dokumentacja procedur', 'Praca ciągła > 72h']:
    add_bullet(doc, e)

add_heading_styled(doc, 'Warunki:', 3)
p = add_body(doc, '')
p.text = ''
for label, val in [('Nakład pracy: ', '120–160 roboczogodzin'), ('Cena ryczałtowa: ', '43 050,00 zł brutto (35 000,00 zł netto + 23% VAT)'), ('Czas realizacji: ', '12 tygodni od odbioru Z1 + dane pilotażowe'), ('Płatności: ', '50% zaliczka + 50% po dostawie')]:
    run = p.add_run(label)
    run.bold = True; run.font.size = Pt(10); run.font.color.rgb = NAVY; run.font.name = 'Segoe UI'
    run = p.add_run(val + '\n')
    run.font.size = Pt(10); run.font.color.rgb = NAVY; run.font.name = 'Segoe UI'

# Z2-B
doc.add_page_break()
add_heading_styled(doc, 'Wariant Z2-B — Profesjonalny (Pełna realizacja Z2)', 2)
add_body(doc, 'Kod oferty: KORDAS/AI/2026/04/Z2-PRO', color=GRAY)

add_heading_styled(doc, 'Zakres prac:', 3)
create_scope_table(doc, [
    ('1', 'Specyfikacja Fazy II', 'Warsztat: scenariusze terenowe, warunki pogodowe, lokalizacje. Aneks.'),
    ('2', 'Pipeline danych', 'Przetwarzanie danych z pilotaży. Anotacja, augmentacja, walidacja.'),
    ('3', 'Retraining iteracyjny', 'Cykliczny trening (min. 3 iteracje) na rosnącym zbiorze.'),
    ('4', 'Eliminacja FP', 'Roślinność, cienie, obiekty ruchome, zmiany oświetlenia. Multi-frame tracking.'),
    ('5', 'Benchmarking', 'Testy porównawcze każdej wersji. Raport z trendami.'),
    ('6', 'IoT ciągła praca', 'Pamięć, termika, watchdog, auto-restart, logging. Cel: > 7 dni stabilnie.'),
    ('7', 'Optymalizacja RPi5', 'Czas reakcji < 1 s przy pracy ciągłej.'),
    ('8', 'Integracja z app', 'Komunikacja z aplikacją użytkownika. Pełna wymiana danych.'),
    ('9', 'Testy finalne', 'Walidacja na instalacji pilotażowej. Raport: dokładność ≥ 75%.'),
    ('10', 'Dokumentacja + IP', 'Pełna dokumentacja + transfer IP + kod + modele + dane + pipeline CI.'),
])

add_body(doc, 'Uwaga dot. progu 75%: Gwarancja warunkowa — wymaga: ukończenia Z1-B, min. 2000 anotowanych klatek z pilotaży, dostępu do instalacji (min. 2 tyg.), terminowych dostaw danych (max 14 dni od zbierania). Niespełnienie → best-effort.', italic=True, color=GRAY)

add_heading_styled(doc, 'Warunki:', 3)
p = add_body(doc, '')
p.text = ''
for label, val in [
    ('Nakład pracy: ', '280–360 roboczogodzin'),
    ('Cena ryczałtowa: ', '98 400,00 zł brutto (80 000,00 zł netto + 23% VAT)'),
    ('Czas realizacji: ', '16 tygodni od Aneksu Specyfikacyjnego Fazy II'),
    ('Płatności: ', '30% zaliczka → 30% po 2. iteracji → 30% po odbiorze → 10% stabilizacja')
]:
    run = p.add_run(label)
    run.bold = True; run.font.size = Pt(10); run.font.color.rgb = NAVY; run.font.name = 'Segoe UI'
    run = p.add_run(val + '\n')
    run.font.size = Pt(10); run.font.color.rgb = NAVY; run.font.name = 'Segoe UI'

# ======================================================================
# PORÓWNANIE
# ======================================================================
doc.add_page_break()
add_heading_styled(doc, 'Porównanie wariantów', 1)

add_heading_styled(doc, 'Zestawienie łączne', 2)
create_comparison_table(doc,
    ['Kombinacja', 'Cena brutto', 'Godziny', 'Uwagi'],
    [
        ['Z1-A + Z2-A (oba bazowe)', '**70 110 zł**', '200–260h', 'PoC + adaptacja dzienna; brak gwarancji progów'],
        ['Z1-B + Z2-A', '**116 850 zł**', '340–440h', 'Solidna faza I + bazowa faza II'],
        ['Z1-B + Z2-B (oba profesjonalne)', '**172 200 zł**', '500–640h', 'Pełna realizacja obu zakresów'],
    ])

# ======================================================================
# WARUNKI OGÓLNE
# ======================================================================
doc.add_page_break()
add_heading_styled(doc, 'Warunki ogólne', 1)

sections_data = [
    ('1. Charakter oferty', 'Niniejsza oferta stanowi szacunek cenowy o charakterze informacyjnym i nie stanowi oferty w rozumieniu art. 66 § 1 Kodeksu Cywilnego. Ostateczny zakres, parametry i warunki finansowe zostaną określone w umowie i Aneksie Specyfikacyjnym.'),
    ('2. Ważność oferty', 'Oferta jest ważna 14 dni kalendarzowych od daty sporządzenia (do 15.04.2026). Po upływie terminu Wykonawca zastrzega prawo do aktualizacji warunków.'),
    ('3. Własność intelektualna', 'Autorskie prawa majątkowe do oprogramowania przeniesione na Zamawiającego po pełnej zapłacie — dotyczy wariantów Profesjonalnych. W Bazowych: kod źródłowy + wagi modelu; pełny transfer IP wymaga wariantu Profesjonalnego.'),
    ('4. Zależności po stronie Zamawiającego', 'Realizacja uzależniona od terminowego dostarczenia: specyfikacji kamery, próbek wideo/zdjęć, specyfikacji prototypu, informacji zwrotnej (5 dni roboczych). Opóźnienia przesuwają harmonogram.'),
    ('5. Gwarancje progów dokładności', 'Progi (65% Z1, 75% Z2) gwarantowane wyłącznie w wariantach Profesjonalnych i pod warunkiem spełnienia wymogów dot. danych. Niespełnienie → tryb best-effort bez odpowiedzialności Wykonawcy.'),
    ('6. Zmiana zakresu', 'Każda zmiana po podpisaniu umowy wymaga formalnego Change Request i aneksu finansowo-terminowego.'),
    ('7. Ograniczenie odpowiedzialności', 'Odpowiedzialność ograniczona do wysokości otrzymanego wynagrodzenia. Wykonawca nie odpowiada za: skuteczność biologiczną, warunki poza specyfikacją, opóźnienia Zamawiającego, jakość danych, kompatybilność z komponentami trzecimi.'),
    ('8. Poufność i RODO', 'Strony zobowiązują się do zachowania poufności. Przy danych z wizerunkami — wymagana umowa DPA.'),
    ('9. Prawo odstąpienia', 'Brak podpisania umowy w ciągu 30 dni od akceptacji → prawo do aktualizacji warunków cenowych.'),
]
for title, body in sections_data:
    add_heading_styled(doc, title, 3)
    add_body(doc, body)

# ---- FOOTER / SIGNATURE ----
doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)
pPr = p._element.get_or_add_pPr()
pBdr = pPr.makeelement(qn('w:pBdr'), {})
top_border = pBdr.makeelement(qn('w:top'), {
    qn('w:val'): 'single', qn('w:sz'): '8', qn('w:color'): '3B82F6', qn('w:space'): '1'
})
pBdr.append(top_border)
pPr.append(pBdr)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run('Sporządził:\n')
run.font.size = Pt(9); run.font.color.rgb = GRAY; run.font.name = 'Segoe UI'
run = p.add_run('Marcin Kordas\n')
run.bold = True; run.font.size = Pt(12); run.font.color.rgb = NAVY; run.font.name = 'Segoe UI'
run = p.add_run('AI Innovation Expert | Kordas Technologies\n')
run.font.size = Pt(10); run.font.color.rgb = GRAY; run.font.name = 'Segoe UI'
run = p.add_run('Kraków, 01.04.2026\n\n')
run.font.size = Pt(9); run.font.color.rgb = GRAY; run.font.name = 'Segoe UI'
run = p.add_run('biuro@kordas.tech | +48 797 252 208')
run.font.size = Pt(9); run.font.color.rgb = BLUE; run.font.name = 'Segoe UI'

# SAVE
out = os.path.join(BASE, "Oferta_KORDAS_AI_2026_04_Antyszpak.docx")
doc.save(out)
print(f"DOCX saved to: {out}")
